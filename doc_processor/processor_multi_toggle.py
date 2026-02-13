from docx import Document
import streamlit as st
import re
import pandas as pd
from io import StringIO
import os
from odf.opendocument import load
from odf.text import P, H, LineBreak


#Global definition of article delimiters and field identifiers
ARTICLE_START = "==Artigo_inicio=="
METADATA_FIELDS = ["#Titulo", "#SubTitulo", "#Autor", "#Data", "#Tag", "#Pag", "#Numero", "#Imagens"]


def extract_text_from_word(file, file_extension):
    """Extract all text from a Word document."""
    if file_extension == ".docx":
        doc = Document(file)
        if output_format == "HTML":
            text = [paragraph_to_html(p) for p in doc.paragraphs]
        else:
            text = [p.text for p in doc.paragraphs]
        return '\n'.join(text)
    elif file_extension == ".odt":
        """
        Extract text from ODT. Collects <text:p> text nodes in document order.
        """
        doc = load(file)
        lines = []
        paragraphs = doc.getElementsByType(P)
        for paragraph in paragraphs:
            text_str = "".join(getattr(n, "data", "") for n in paragraph.childNodes).strip()
            if text_str:
                lines.append(text_str)


    return "\n".join(lines)

def article_split(text):
    """Split articles and extract body and metadata.
    This version treats lines like 'Titulo: valor' (without leading '#') as metadata.
    It also accepts legacy '#Titulo: valor' forms.
    """
    raw_articles = text.split(ARTICLE_START)
    articles = []

    # loops articles
    for raw_article in raw_articles:
        raw_article = raw_article.strip()
        if not raw_article:
            continue  # skip empty articles

        lines = [line.strip() for line in raw_article.splitlines() if line.strip()]

        metadata = {}
        body_lines = []

        # loops lines to find metadata fields to add them to the metadata dictionary
        for line in lines:
            is_metadata = any(line.startswith(f"{field}:") for field in METADATA_FIELDS)
            if is_metadata:
                key, value = line.split(":", 1)
                field = key.strip()
                metadata[field] = value.strip()
            else:
                body_lines.append(line)

        # join body lines and try to split Rodape using the appropriate marker
        body_text = "\n".join(body_lines).strip()
        try:
            split_marker = "<p>#Rodape:" if output_format == "HTML" else "#Rodape:"
            if split_marker in body_text:
                body_part, rodape_part = body_text.split(split_marker, 1)
                metadata["BODY"] = body_part.strip()
                metadata["Rodape"] = ("<p>" + rodape_part) if output_format == "HTML" else rodape_part.strip()
            else:
                metadata["BODY"] = body_text
        except Exception:
            metadata["BODY"] = body_text

        # Fill missing metadata keys with empty string
        for field in METADATA_FIELDS:
            if field not in metadata:
                metadata[field] = ""

        articles.append(metadata)

    return articles


def paragraph_to_html(paragraph):
    html = ""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}

    # Process paragraph children (mix of runs and hyperlinks)
    for child in paragraph._element:
        # Handle hyperlink elements that contain runs
        if child.tag.endswith('hyperlink'):
            rId = child.get('{%s}id' % ns['r'])
            if rId:
                try:
                    url = paragraph.part.rels[rId].target_ref
                    # Extract text from all runs inside the hyperlink
                    link_text = "".join(run.text for run in child.findall('.//w:r', ns) if run.text)
                    if link_text:
                        html += f'<a href="{url}">{link_text}</a>'
                except (KeyError, AttributeError):
                    pass
        # Handle regular runs
        elif child.tag.endswith('r'):
            run_elem = child
            text = ""
            for t_elem in run_elem.findall('.//w:t', ns):
                text += t_elem.text or ""

            if text:
                if run_elem.find('.//w:b', ns) is not None:
                    text = f"<b>{text}</b>"
                if run_elem.find('.//w:i', ns) is not None:
                    text = f"<i>{text}</i>"
                if run_elem.find('.//w:u', ns) is not None:
                    text = f"<u>{text}</u>"
                html += text

    # treat paragraphs that are article start or metadata lines without leading '#'
    if html.startswith(ARTICLE_START) or any(html.startswith(f"{field}:") for field in METADATA_FIELDS):
        return html
    else:
        return f"<p>{html}</p>"


if __name__ == "__main__":
    # --- Streamlit app ---
    st.title("📄 Separador Automático de artigos")
    st.write("Faça upload do documento .docx ou .odt para separar automaticamente artigos")

    output_format = st.radio("Escolha o formato de saída:", ("Texto", "HTML"), key="output_format")

    toggle_option = st.radio("Modo de delimitadores:", ("Padrão", "Personalizado"), key="toggle_option")

    if toggle_option == "Personalizado":
        st.subheader("Configuração Personalizada")

        article_start_input = st.text_input(
            "Delimitador de Artigo",
            value=ARTICLE_START,
            placeholder="Ex: ==Artigo_inicio==",
            help="Texto que marca o início de um novo artigo"
        )
        ARTICLE_START = article_start_input

        metadata_fields_input = st.text_area(
            "Campos de Metadados",
            value=", ".join(METADATA_FIELDS),
            placeholder="Ex: #Titulo, #SubTitulo, #Autor, #Data, #Tag, #Pag, #Numero, #Imagens",
            height=100,
            help="Separe os campos por vírgula"
        )
        METADATA_FIELDS = [field.strip() for field in metadata_fields_input.split(",")]
    else:
        st.subheader("Configuração Padrão")
        st.text_input(
            "Delimitador de Artigo",
            value=ARTICLE_START,
            disabled=True
        )
        st.text_area(
            "Campos de Metadados",
            value=", ".join(METADATA_FIELDS),
            height=100,
            disabled=True
        )
    uploaded_file = st.file_uploader("Escolha um documento .docx", type=["docx", "odt"])

    if uploaded_file:
        filename, file_extension = os.path.splitext(uploaded_file.name)

        text = extract_text_from_word(uploaded_file, file_extension)
        articles = article_split(text)

        df = pd.DataFrame(articles)

        # Display the DataFrame as a CSV table
        st.success(f"✅ encontrados {len(articles)} artigos.")
        st.dataframe(df)

        # Generate CSV for download
        csv_buffer = StringIO()
        df.to_csv(csv_buffer, index=False)
        csv_data = csv_buffer.getvalue()

        st.download_button(
            label="💾 Download CSV",
            data=csv_data,
            file_name="articles.csv",
            mime="text/csv"
        )

        for i, article in enumerate(articles, 1):
            st.subheader(f"Artigo {i}")
            st.text_area(article.get("Titulo", ""), article.get("BODY", ""), height=200)

            if "Rodape" in article.keys() and article.get("Rodape"):
                st.text_area("Rodapé", article.get("Rodape", ""), height=60)

    else:
        st.info("Escolha um documento .docx para começar.")