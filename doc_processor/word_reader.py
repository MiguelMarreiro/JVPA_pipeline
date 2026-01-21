from docx import Document
from pathlib import Path
from typing import Union


def extract_text_from_word(file_path: Union[str, Path]) -> str:
    """
    Read and extract all text from a Word (.docx) file.
    
    Args:
        file_path (str): Path to the Word document file (.docx)
        
    Returns:
        str: Extracted text from the document
        
    Raises:
        FileNotFoundError: If the file does not exist
        ValueError: If the file is not a valid .docx file
        
    Example:
        text = extract_text_from_word('document.docx')
        print(text)
    """
    file_path = Path(file_path)
    
    # Validate file exists
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Validate file extension
    if file_path.suffix.lower() != '.docx':
        raise ValueError(f"File must be a .docx file, got: {file_path.suffix}")
    
    try:
        # Load the Word document
        doc = Document(str(file_path))
        
        # Extract text from all paragraphs
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return '\n'.join(text)
    
    except Exception as e:
        raise ValueError(f"Error reading Word document: {str(e)}")


def extract_text_from_word_with_metadata(file_path: Union[str, Path]) -> dict:
    """
    Read and extract text from a Word file along with metadata.
    
    Args:
        file_path (str): Path to the Word document file (.docx)
        
    Returns:
        dict: Dictionary containing:
            - 'text': Extracted text from the document
            - 'paragraph_count': Number of paragraphs
            - 'table_count': Number of tables
            - 'file_name': Name of the file
            
    Raises:
        FileNotFoundError: If the file does not exist
        ValueError: If the file is not a valid .docx file
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if file_path.suffix.lower() != '.docx':
        raise ValueError(f"File must be a .docx file, got: {file_path.suffix}")
    
    try:
        doc = Document(str(file_path))
        
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return {
            'text': '\n'.join(text),
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'file_name': file_path.name
        }
    
    except Exception as e:
        raise ValueError(f"Error reading Word document: {str(e)}")


if __name__ == '__main__':
    # Example usage
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            text = extract_text_from_word(file_path)
            print("Extracted Text:")
            print(text)
        except (FileNotFoundError, ValueError) as e:
            print(f"Error: {e}")
    else:
        print("Usage: python word_reader.py <path_to_word_file>")
